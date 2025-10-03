"""
Text Parser for Cultural-Specific Information Extraction

This module handles Markdown (.md) and Word (.docx) files to extract cultural-specific
nutritional information including:
- Traditional foods and recipes
- Cultural dietary practices
- Regional food preferences
- Religious dietary restrictions
- Festival and ceremonial foods
- Cultural cooking methods
- Traditional medicine foods
"""

import logging
import re
import json
from typing import Dict, List, Any, Optional, Set
from pathlib import Path
import markdown
from dataclasses import dataclass, asdict

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX parsing will be limited.")


@dataclass
class CulturalData:
    """Data class for structured cultural nutritional information"""

    culture_name: str = ""
    region: str = ""
    traditional_foods: List[str] = None
    dietary_practices: List[str] = None
    religious_restrictions: List[str] = None
    festive_foods: Dict[str, List[str]] = None  # festival/occasion -> foods
    cooking_methods: List[str] = None
    medicinal_foods: Dict[str, str] = None  # food -> medicinal use
    seasonal_foods: Dict[str, List[str]] = None  # season -> foods
    food_combinations: List[str] = None
    cultural_beliefs: List[str] = None
    prohibited_foods: List[str] = None
    preferred_flavors: List[str] = None
    meal_patterns: Dict[str, str] = None  # meal_type -> description
    age_specific_foods: Dict[str, List[str]] = None  # age_group -> foods
    gender_specific_foods: Dict[str, List[str]] = None  # gender -> foods
    source_document: str = ""
    extraction_confidence: float = 0.0

    def __post_init__(self):
        """Initialize mutable default values"""
        if self.traditional_foods is None:
            self.traditional_foods = []
        if self.dietary_practices is None:
            self.dietary_practices = []
        if self.religious_restrictions is None:
            self.religious_restrictions = []
        if self.festive_foods is None:
            self.festive_foods = {}
        if self.cooking_methods is None:
            self.cooking_methods = []
        if self.medicinal_foods is None:
            self.medicinal_foods = {}
        if self.seasonal_foods is None:
            self.seasonal_foods = {}
        if self.food_combinations is None:
            self.food_combinations = []
        if self.cultural_beliefs is None:
            self.cultural_beliefs = []
        if self.prohibited_foods is None:
            self.prohibited_foods = []
        if self.preferred_flavors is None:
            self.preferred_flavors = []
        if self.meal_patterns is None:
            self.meal_patterns = {}
        if self.age_specific_foods is None:
            self.age_specific_foods = {}
        if self.gender_specific_foods is None:
            self.gender_specific_foods = {}


class TextCulturalParser:
    """Main class for parsing cultural information from text files"""

    def __init__(self):
        self.cultural_patterns = self._initialize_patterns()
        self.culture_keywords = self._initialize_culture_keywords()
        self.food_keywords = self._initialize_food_keywords()

    def _initialize_patterns(self) -> Dict[str, re.Pattern]:
        """Initialize regex patterns for cultural data extraction"""
        patterns = {
            # Culture identification
            "culture_name": re.compile(
                r"(?:culture|tradition|community)[:\s]*([A-Za-z\s,]+)", re.IGNORECASE
            ),
            "region": re.compile(
                r"(?:region|area|country|state)[:\s]*([A-Za-z\s,]+)", re.IGNORECASE
            ),
            # Traditional foods
            "traditional_foods": re.compile(
                r"(?:traditional|staple|native|indigenous)\s+(?:food|dish|meal)[s]?[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            "regional_dishes": re.compile(
                r"(?:regional|local|popular)\s+(?:dish|food|cuisine)[s]?[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            # Dietary practices
            "dietary_practice": re.compile(
                r"(?:eat|consume|practice|custom)[s]?[:\s]*([^.!?]*)", re.IGNORECASE
            ),
            "eating_pattern": re.compile(
                r"(?:meal|eating)\s+(?:pattern|habit|time)[s]?[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            # Religious restrictions
            "religious_restriction": re.compile(
                r"(?:forbidden|prohibited|taboo|restriction|avoid)[s]?[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            "halal_kosher": re.compile(
                r"(?:halal|kosher|permissible)[:\s]*([^.!?]*)", re.IGNORECASE
            ),
            # Festivals and occasions
            "festival_food": re.compile(
                r"(?:festival|celebration|ceremony|occasion)[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            "special_occasion": re.compile(
                r"(?:wedding|birth|death|harvest)[:\s]*([^.!?]*)", re.IGNORECASE
            ),
            # Cooking methods
            "cooking_method": re.compile(
                r"(?:cook|prepare|method|technique)[s]?[:\s]*([^.!?]*)", re.IGNORECASE
            ),
            "preparation": re.compile(
                r"(?:preparation|cooking|processing)[:\s]*([^.!?]*)", re.IGNORECASE
            ),
            # Medicinal foods
            "medicinal_food": re.compile(
                r"(?:medicinal|healing|therapeutic|remedy)[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            "health_benefit": re.compile(
                r"(?:benefit|property|effect|use)[s]?[:\s]*([^.!?]*)", re.IGNORECASE
            ),
            # Seasonal foods
            "seasonal": re.compile(
                r"(?:season|seasonal|summer|winter|spring|autumn|monsoon)[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            # Age and gender specific
            "age_specific": re.compile(
                r"(?:children|adult|elderly|infant|toddler|teenager)[s]?[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            "gender_specific": re.compile(
                r"(?:women|men|male|female|pregnant|lactating)[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
            # Food combinations
            "combination": re.compile(
                r"(?:combine|together|with|paired|served)[:\s]*([^.!?]*)", re.IGNORECASE
            ),
            # Flavors and preferences
            "flavor": re.compile(
                r"(?:flavor|taste|spicy|sweet|sour|bitter|salty)[s]?[:\s]*([^.!?]*)",
                re.IGNORECASE,
            ),
        }
        return patterns

    def _initialize_culture_keywords(self) -> Set[str]:
        """Initialize keywords that indicate cultural content"""
        return {
            "indian",
            "chinese",
            "japanese",
            "thai",
            "mexican",
            "italian",
            "french",
            "mediterranean",
            "middle eastern",
            "african",
            "latin american",
            "hindu",
            "muslim",
            "buddhist",
            "christian",
            "jewish",
            "sikh",
            "ayurvedic",
            "traditional",
            "cultural",
            "regional",
            "ethnic",
            "festival",
            "ceremony",
            "ritual",
            "celebration",
            "custom",
            "tribe",
            "community",
            "village",
            "rural",
            "urban",
        }

    def _initialize_food_keywords(self) -> Set[str]:
        """Initialize food-related keywords"""
        return {
            "rice",
            "wheat",
            "bread",
            "curry",
            "dal",
            "spice",
            "herb",
            "vegetable",
            "fruit",
            "meat",
            "fish",
            "dairy",
            "oil",
            "tea",
            "coffee",
            "milk",
            "yogurt",
            "cheese",
            "butter",
            "soup",
            "stew",
            "sauce",
            "pickle",
            "fermented",
            "preserved",
        }

    def parse_markdown(self, file_path: str) -> List[CulturalData]:
        """
        Parse a Markdown file and extract cultural information

        Args:
            file_path: Path to the Markdown file

        Returns:
            List of CulturalData objects containing extracted information
        """
        logger.info(f"Parsing Markdown file: {file_path}")

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Convert markdown to HTML and extract text
            html_content = markdown.markdown(content)

            # Also work with raw markdown for better pattern matching
            cultural_data = self._extract_from_text(content, file_path)

            # Extract structured data from markdown headers and lists
            structured_data = self._extract_from_markdown_structure(content, file_path)
            cultural_data.extend(structured_data)

            logger.info(
                f"Extracted {len(cultural_data)} cultural entries from Markdown"
            )
            return cultural_data

        except Exception as e:
            logger.error(f"Error parsing Markdown file {file_path}: {str(e)}")
            return []

    def parse_docx(self, file_path: str) -> List[CulturalData]:
        """
        Parse a DOCX file and extract cultural information

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of CulturalData objects containing extracted information
        """
        logger.info(f"Parsing DOCX file: {file_path}")

        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Cannot parse DOCX files.")
            return []

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            full_text = ""
            sections = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text += text + "\n"

                    # Check if this is a heading (based on style or formatting)
                    if paragraph.style.name.startswith("Heading") or len(text) < 100:
                        sections.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        full_text += row_text + "\n"

            # Process the extracted text
            cultural_data = self._extract_from_text(full_text, file_path)

            logger.info(f"Extracted {len(cultural_data)} cultural entries from DOCX")
            return cultural_data

        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            return []

    def _extract_from_text(self, text: str, source_file: str) -> List[CulturalData]:
        """Extract cultural information from plain text"""
        logger.info("Extracting cultural data from text content")

        # Split text into sections
        sections = self._split_into_cultural_sections(text)
        cultural_data = []

        for section in sections:
            if self._contains_cultural_content(section):
                data = self._extract_cultural_values(section, source_file)
                if data and self._is_valid_cultural_data(data):
                    cultural_data.append(data)

        return cultural_data

    def _extract_from_markdown_structure(
        self, content: str, source_file: str
    ) -> List[CulturalData]:
        """Extract cultural information from markdown structure (headers, lists, etc.)"""
        cultural_data = []

        # Split by headers
        sections = re.split(r"^#+\s+(.+)$", content, flags=re.MULTILINE)

        current_culture = CulturalData(source_document=source_file)

        for i in range(
            1, len(sections), 2
        ):  # Skip empty parts, take header+content pairs
            if i + 1 < len(sections):
                header = sections[i].strip()
                section_content = sections[i + 1].strip()

                # Process based on header content
                if any(
                    keyword in header.lower()
                    for keyword in ["culture", "tradition", "community"]
                ):
                    current_culture.culture_name = header
                elif any(
                    keyword in header.lower() for keyword in ["food", "dish", "cuisine"]
                ):
                    foods = self._extract_list_items(section_content)
                    current_culture.traditional_foods.extend(foods)
                elif any(
                    keyword in header.lower() for keyword in ["festival", "celebration"]
                ):
                    festival_foods = self._extract_festival_foods(section_content)
                    current_culture.festive_foods.update(festival_foods)
                elif any(
                    keyword in header.lower() for keyword in ["cooking", "preparation"]
                ):
                    methods = self._extract_list_items(section_content)
                    current_culture.cooking_methods.extend(methods)
                elif any(
                    keyword in header.lower() for keyword in ["medicinal", "healing"]
                ):
                    medicinal = self._extract_medicinal_foods(section_content)
                    current_culture.medicinal_foods.update(medicinal)

        if self._is_valid_cultural_data(current_culture):
            cultural_data.append(current_culture)

        return cultural_data

    def _split_into_cultural_sections(self, text: str) -> List[str]:
        """Split text into sections that might contain cultural information"""
        patterns = [
            r"\n\s*\n",  # Double newlines
            r"(?=\b(?:Culture|Tradition|Community|Region)\b)",
            r"(?=\b(?:Festival|Celebration|Ceremony)\b)",
            r"(?=\b(?:Food|Cuisine|Dish|Recipe)\b)",
        ]

        sections = [text]
        for pattern in patterns:
            new_sections = []
            for section in sections:
                new_sections.extend(re.split(pattern, section, flags=re.IGNORECASE))
            sections = [s.strip() for s in new_sections if s.strip() and len(s) > 50]

        return sections

    def _contains_cultural_content(self, text: str) -> bool:
        """Check if text contains cultural content"""
        text_lower = text.lower()

        # Check for culture keywords
        culture_score = sum(
            1 for keyword in self.culture_keywords if keyword in text_lower
        )
        food_score = sum(1 for keyword in self.food_keywords if keyword in text_lower)

        return culture_score >= 1 and food_score >= 1

    def _extract_cultural_values(
        self, text: str, source_file: str
    ) -> Optional[CulturalData]:
        """Extract cultural values from a text section"""
        data = CulturalData(source_document=source_file)
        extraction_count = 0

        # Extract each type of cultural information
        for info_type, pattern in self.cultural_patterns.items():
            matches = pattern.findall(text)

            if matches:
                extraction_count += 1

                # Process different types of information
                if info_type == "culture_name":
                    data.culture_name = matches[0].strip()
                elif info_type == "region":
                    data.region = matches[0].strip()
                elif info_type in ["traditional_foods", "regional_dishes"]:
                    foods = self._parse_food_list(matches[0])
                    data.traditional_foods.extend(foods)
                elif info_type in ["dietary_practice", "eating_pattern"]:
                    practices = self._parse_practice_list(matches[0])
                    data.dietary_practices.extend(practices)
                elif info_type in ["religious_restriction", "halal_kosher"]:
                    restrictions = self._parse_restriction_list(matches[0])
                    data.religious_restrictions.extend(restrictions)
                elif info_type in ["festival_food", "special_occasion"]:
                    festival_data = self._parse_festival_foods(matches[0])
                    data.festive_foods.update(festival_data)
                elif info_type in ["cooking_method", "preparation"]:
                    methods = self._parse_method_list(matches[0])
                    data.cooking_methods.extend(methods)
                elif info_type in ["medicinal_food", "health_benefit"]:
                    medicinal = self._parse_medicinal_foods(matches[0])
                    data.medicinal_foods.update(medicinal)
                elif info_type == "flavor":
                    flavors = self._parse_flavor_list(matches[0])
                    data.preferred_flavors.extend(flavors)

        # Calculate extraction confidence
        data.extraction_confidence = min(extraction_count / 8.0, 1.0)

        return data if extraction_count > 0 else None

    def _extract_list_items(self, content: str) -> List[str]:
        """Extract items from markdown lists or comma-separated text"""
        items = []

        # Try markdown list format
        list_items = re.findall(r"^[\s]*[-*+]\s+(.+)$", content, re.MULTILINE)
        items.extend([item.strip() for item in list_items])

        # Try numbered list
        numbered_items = re.findall(r"^\s*\d+\.\s+(.+)$", content, re.MULTILINE)
        items.extend([item.strip() for item in numbered_items])

        # Try comma-separated
        if not items:
            comma_items = [item.strip() for item in content.split(",")]
            items.extend([item for item in comma_items if len(item) > 2])

        return items[:10]  # Limit to avoid noise

    def _extract_festival_foods(self, content: str) -> Dict[str, List[str]]:
        """Extract festival-food mappings"""
        festival_foods = {}

        # Look for patterns like "Diwali: sweets, nuts" or "Christmas - cookies, cake"
        patterns = [
            r"([A-Za-z\s]+)[:]\s*([^.!?\n]+)",
            r"([A-Za-z\s]+)[-–]\s*([^.!?\n]+)",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, content)
            for festival, foods in matches:
                festival = festival.strip()
                food_list = [f.strip() for f in foods.split(",")]
                festival_foods[festival] = food_list

        return festival_foods

    def _extract_medicinal_foods(self, content: str) -> Dict[str, str]:
        """Extract medicinal food-use mappings"""
        medicinal_foods = {}

        # Look for patterns like "ginger: digestive aid" or "turmeric - anti-inflammatory"
        patterns = [
            r"([A-Za-z\s]+)[:]\s*([^.!?\n]+)",
            r"([A-Za-z\s]+)[-–]\s*([^.!?\n]+)",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, content)
            for food, use in matches:
                medicinal_foods[food.strip()] = use.strip()

        return medicinal_foods

    def _parse_food_list(self, text: str) -> List[str]:
        """Parse food items from text"""
        foods = []

        # Split by common delimiters
        items = re.split(r"[,;]\s*", text)
        for item in items:
            item = item.strip()
            if item and len(item) > 1:
                foods.append(item)

        return foods

    def _parse_practice_list(self, text: str) -> List[str]:
        """Parse dietary practices from text"""
        return self._parse_food_list(text)  # Similar parsing logic

    def _parse_restriction_list(self, text: str) -> List[str]:
        """Parse dietary restrictions from text"""
        return self._parse_food_list(text)  # Similar parsing logic

    def _parse_festival_foods(self, text: str) -> Dict[str, List[str]]:
        """Parse festival foods from text"""
        return self._extract_festival_foods(text)

    def _parse_method_list(self, text: str) -> List[str]:
        """Parse cooking methods from text"""
        return self._parse_food_list(text)  # Similar parsing logic

    def _parse_medicinal_foods(self, text: str) -> Dict[str, str]:
        """Parse medicinal foods from text"""
        return self._extract_medicinal_foods(text)

    def _parse_flavor_list(self, text: str) -> List[str]:
        """Parse flavor preferences from text"""
        return self._parse_food_list(text)  # Similar parsing logic

    def _is_valid_cultural_data(self, data: CulturalData) -> bool:
        """Check if the extracted data contains meaningful cultural information"""
        has_foods = bool(data.traditional_foods)
        has_practices = bool(data.dietary_practices)
        has_cultural_info = bool(data.culture_name or data.region)
        has_festivals = bool(data.festive_foods)
        has_methods = bool(data.cooking_methods)

        return (
            has_foods or has_practices or has_festivals or has_methods
        ) and has_cultural_info

    def save_results(self, cultural_data: List[CulturalData], output_path: str) -> bool:
        """
        Save extracted cultural data to JSON file

        Args:
            cultural_data: List of CulturalData objects
            output_path: Path where to save the results

        Returns:
            bool: True if saved successfully, False otherwise
        """
        try:
            # Convert dataclass objects to dictionaries
            data_dicts = [asdict(data) for data in cultural_data]

            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(data_dicts, f, indent=2, ensure_ascii=False)

            logger.info(
                f"Successfully saved {len(cultural_data)} entries to {output_path}"
            )
            return True

        except Exception as e:
            logger.error(f"Error saving results to {output_path}: {str(e)}")
            return False


# Example usage and testing
if __name__ == "__main__":
    # Example usage
    parser = TextCulturalParser()

    # Test with sample files
    sample_md = "sample_cultural_info.md"
    sample_docx = "sample_cultural_info.docx"

    all_results = []

    # Parse Markdown file
    if Path(sample_md).exists():
        md_results = parser.parse_markdown(sample_md)
        all_results.extend(md_results)
        print(f"Extracted {len(md_results)} cultural entries from Markdown")

    # Parse DOCX file
    if Path(sample_docx).exists():
        docx_results = parser.parse_docx(sample_docx)
        all_results.extend(docx_results)
        print(f"Extracted {len(docx_results)} cultural entries from DOCX")

    if all_results:
        print(f"\nTotal extracted: {len(all_results)} cultural entries")
        for i, data in enumerate(all_results, 1):
            print(f"\n--- Entry {i} ---")
            print(f"Culture: {data.culture_name}")
            print(f"Region: {data.region}")
            print(f"Traditional Foods: {data.traditional_foods[:3]}")  # Show first 3
            print(f"Dietary Practices: {data.dietary_practices[:3]}")
            print(f"Festival Foods: {list(data.festive_foods.keys())[:3]}")
            print(f"Confidence: {data.extraction_confidence:.2f}")

        # Save results
        output_file = "extracted_cultural_data.json"
        parser.save_results(all_results, output_file)
    else:
        print("No sample files found. Please provide valid MD or DOCX files.")
